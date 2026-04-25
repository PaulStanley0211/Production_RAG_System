"""DOCX extractor — walks paragraphs and tables in document order.

Word documents interleave paragraphs and tables. Walking element-by-element
in body order preserves the natural reading sequence. Tables are flattened
to tab-separated rows. Heading styles are detected and preserved inline.
"""

import logging
from pathlib import Path

from docx import Document as DocxDocument
from docx.oxml.ns import qn

from app.models import Document
from pipeline.extractors.base import Extractor

log = logging.getLogger(__name__)


class DOCXExtractor(Extractor):
    """Extracts text from .docx files in original document order."""

    def can_handle(self, path: Path) -> bool:
        return path.suffix.lower() == ".docx"

    def extract(self, path: Path) -> Document:
        doc = DocxDocument(str(path))

        parts: list[str] = []
        heading_count = 0
        table_count = 0

        # Walk the document body in original order. Each child is either
        # a <w:p> (paragraph) or <w:tbl> (table) — we map back to the
        # python-docx wrapper objects for clean access.
        for child in doc.element.body.iterchildren():
            if child.tag == qn("w:p"):
                # Find the matching paragraph wrapper
                para = next(
                    (p for p in doc.paragraphs if p._element is child),
                    None,
                )
                if para is None or not para.text.strip():
                    continue

                # Detect heading style and prefix with markdown-style hashes
                style_name = (para.style.name or "").lower()
                if style_name.startswith("heading"):
                    heading_count += 1
                    # "heading 1" → "# ", "heading 2" → "## ", etc.
                    try:
                        level = int(style_name.split()[-1])
                        prefix = "#" * min(level, 6) + " "
                    except (ValueError, IndexError):
                        prefix = "# "
                    parts.append(prefix + para.text.strip())
                else:
                    parts.append(para.text.strip())

            elif child.tag == qn("w:tbl"):
                # Find the matching table wrapper
                table = next(
                    (t for t in doc.tables if t._element is child),
                    None,
                )
                if table is None:
                    continue
                table_count += 1
                parts.append(self._table_to_text(table))

        content = "\n\n".join(parts)

        return Document(
            id=str(path.resolve()),
            source=path.name,
            content=content,
            metadata={
                "format": "docx",
                "heading_count": heading_count,
                "table_count": table_count,
                "size_bytes": path.stat().st_size,
            },
        )

    def _table_to_text(self, table) -> str:
        """Flatten a docx table to tab-separated rows."""
        rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):  # skip entirely empty rows
                rows.append("\t".join(cells))
        return "\n".join(rows)